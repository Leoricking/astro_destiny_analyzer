"""
Language-aware Compatibility Report exporters.

Supports Markdown, HTML, Word and PDF. Calculation models stay canonical;
localization is applied only during rendering.
"""
from __future__ import annotations

from datetime import datetime
from html import escape
from io import BytesIO
from typing import TYPE_CHECKING, Dict, List
import re as _re

if TYPE_CHECKING:
    from compatibility.models import CompatibilityReport

_LANGS = ("zh-TW", "en", "th", "ja", "es", "ar")
_META = {
    "zh-TW": ("zh-TW", "ltr"), "en": ("en", "ltr"), "th": ("th", "ltr"),
    "ja": ("ja", "ltr"), "es": ("es", "ltr"), "ar": ("ar", "rtl"),
}

_TXT: Dict[str, Dict[str, str]] = {
"zh-TW": {
"title":"關係合盤精準分析報告","relationship_type":"關係類型","person_a":"A 方","person_b":"B 方",
"created":"產生時間","overall":"關係總分","precision":"精準分析摘要","dimension":"各維度評分",
"emotional":"情感共鳴","communication":"溝通契合","attraction":"吸引力","stability":"穩定性",
"growth":"成長潛力","conflict":"衝突強度","collaboration":"協作效能","strengths":"關係優勢",
"challenges":"關係挑戰","actions":"優先改善順序","evidence":"分數形成依據",
"limits":"資料限制與可信度","practice":"30 天關係練習","long_term":"長期關係建議",
"disclaimer":"免責聲明","high":"高","medium":"中","low":"低",
},
"en": {
"title":"Precision Compatibility Analysis Report","relationship_type":"Relationship Type","person_a":"Person A","person_b":"Person B",
"created":"Created","overall":"Overall Compatibility Score","precision":"Precision Summary","dimension":"Dimension Scores",
"emotional":"Emotional Resonance","communication":"Communication","attraction":"Attraction","stability":"Stability",
"growth":"Growth Potential","conflict":"Conflict Intensity","collaboration":"Collaboration","strengths":"Relationship Strengths",
"challenges":"Relationship Challenges","actions":"Priority Actions","evidence":"Score Evidence",
"limits":"Data Limitations & Confidence","practice":"30-Day Relationship Practice","long_term":"Long-Term Guidance",
"disclaimer":"Disclaimer","high":"High","medium":"Medium","low":"Low",
},
"ja": {
"title":"相性・関係性 精密分析レポート","relationship_type":"関係タイプ","person_a":"Aさん","person_b":"Bさん",
"created":"作成日時","overall":"総合相性スコア","precision":"精密分析サマリー","dimension":"各次元のスコア",
"emotional":"感情的共鳴","communication":"コミュニケーション","attraction":"引力・親密さ","stability":"安定性",
"growth":"成長可能性","conflict":"葛藤強度","collaboration":"協力効率","strengths":"関係の強み",
"challenges":"関係の課題","actions":"優先アクション","evidence":"スコアの根拠",
"limits":"データ制限と信頼度","practice":"30日間の関係練習","long_term":"長期的な提案",
"disclaimer":"免責事項","high":"高","medium":"中","low":"低",
},
"th": {
"title":"รายงานวิเคราะห์ความเข้ากันได้เชิงลึก","relationship_type":"ประเภทความสัมพันธ์","person_a":"บุคคล A","person_b":"บุคคล B",
"created":"สร้างเมื่อ","overall":"คะแนนความเข้ากันได้โดยรวม","precision":"สรุปการวิเคราะห์เชิงลึก","dimension":"คะแนนรายมิติ",
"emotional":"ความสอดคล้องทางอารมณ์","communication":"การสื่อสาร","attraction":"แรงดึงดูด","stability":"ความมั่นคง",
"growth":"ศักยภาพการเติบโต","conflict":"ความเข้มข้นของความขัดแย้ง","collaboration":"การร่วมมือ","strengths":"จุดแข็งของความสัมพันธ์",
"challenges":"ความท้าทาย","actions":"ลำดับการปรับปรุง","evidence":"หลักฐานของคะแนน",
"limits":"ข้อจำกัดของข้อมูลและความเชื่อมั่น","practice":"แบบฝึกหัดความสัมพันธ์ 30 วัน","long_term":"คำแนะนำระยะยาว",
"disclaimer":"ข้อจำกัดความรับผิดชอบ","high":"สูง","medium":"ปานกลาง","low":"ต่ำ",
},
"es": {
"title":"Informe de compatibilidad de precisión","relationship_type":"Tipo de relación","person_a":"Persona A","person_b":"Persona B",
"created":"Creado","overall":"Puntuación global de compatibilidad","precision":"Resumen de precisión","dimension":"Puntuaciones por dimensión",
"emotional":"Resonancia emocional","communication":"Comunicación","attraction":"Atracción","stability":"Estabilidad",
"growth":"Potencial de crecimiento","conflict":"Intensidad del conflicto","collaboration":"Colaboración","strengths":"Fortalezas de la relación",
"challenges":"Desafíos de la relación","actions":"Acciones prioritarias","evidence":"Evidencia de puntuación",
"limits":"Limitaciones y confianza","practice":"Práctica de relación de 30 días","long_term":"Orientación a largo plazo",
"disclaimer":"Descargo de responsabilidad","high":"Alta","medium":"Media","low":"Baja",
},
"ar": {
"title":"تقرير دقيق لتحليل التوافق","relationship_type":"نوع العلاقة","person_a":"الشخص أ","person_b":"الشخص ب",
"created":"تاريخ الإنشاء","overall":"درجة التوافق العامة","precision":"ملخص التحليل الدقيق","dimension":"درجات الأبعاد",
"emotional":"الانسجام العاطفي","communication":"التواصل","attraction":"الانجذاب","stability":"الاستقرار",
"growth":"إمكانات النمو","conflict":"شدة الصراع","collaboration":"التعاون","strengths":"نقاط قوة العلاقة",
"challenges":"تحديات العلاقة","actions":"الإجراءات ذات الأولوية","evidence":"أدلة الدرجة",
"limits":"قيود البيانات ومستوى الثقة","practice":"تمرين علاقة لمدة 30 يومًا","long_term":"إرشاد طويل المدى",
"disclaimer":"إخلاء المسؤولية","high":"مرتفعة","medium":"متوسطة","low":"منخفضة",
},
}

_GENERIC = {
"en": {
"summary":"Overall score {overall}/100. The strongest dimension is {best} ({best_score}/100); the main development area is {weakest} ({weak_score}/100). Conflict intensity is {conflict}/100.",
"strength":"{name} is currently a relative strength at {score}/100. Protect it through consistent behavior rather than assumptions.",
"challenge":"{name} is the lowest-scoring area at {score}/100. Discuss concrete expectations, timing, and boundaries.",
"action_comm":"Use a weekly 20-minute check-in: one observation, one feeling, one request, and one agreed action.",
"action_emotion":"Before problem-solving, name the emotion and confirm what kind of support is wanted.",
"action_conflict":"When tension rises, pause for at least 20 minutes and agree on a specific time to resume the conversation.",
"action_stability":"Clarify responsibilities, money, time, exclusivity, and decision rights in writing.",
"limit":"Scores depend on the birth data provided. Missing birth time or coordinates reduces the precision of Ascendant, houses, Zi Wei timing, and composite angles.",
"disclaimer":"This report supports reflection and communication. It does not determine whether a relationship must succeed or fail and is not medical, legal, or financial advice.",
},
"ja": {
"summary":"総合スコアは{overall}/100です。最も強い次元は{best}（{best_score}/100）、主な改善領域は{weakest}（{weak_score}/100）です。葛藤強度は{conflict}/100です。",
"strength":"{name}は{score}/100で、現在の相対的な強みです。思い込みではなく、継続的な行動で守りましょう。",
"challenge":"{name}は{score}/100で最も低い領域です。期待、タイミング、境界線を具体的に話し合ってください。",
"action_comm":"週1回20分の確認時間を設け、「観察・感情・要望・合意した行動」を一つずつ共有します。",
"action_emotion":"解決策を出す前に感情を言語化し、どのような支援を望むか確認します。",
"action_conflict":"緊張が高まったら20分以上中断し、再開時刻を具体的に決めます。",
"action_stability":"責任、金銭、時間、排他性、意思決定権を明確にします。",
"limit":"結果は入力された出生情報に依存します。出生時刻や座標がない場合、アセンダント、ハウス、紫微の時刻要素、Composite軸の精度が下がります。",
"disclaimer":"本レポートは理解と対話のための参考です。関係の成功・失敗を決定するものではなく、医療・法律・金融の助言ではありません。",
},
"th": {
"summary":"คะแนนรวม {overall}/100 มิติที่แข็งแรงที่สุดคือ {best} ({best_score}/100) ส่วนที่ควรพัฒนามากที่สุดคือ {weakest} ({weak_score}/100) และความเข้มข้นของความขัดแย้งคือ {conflict}/100",
"strength":"{name} เป็นจุดแข็งสัมพัทธ์ที่ {score}/100 ควรรักษาด้วยพฤติกรรมที่สม่ำเสมอ ไม่ใช่การคาดเดา",
"challenge":"{name} เป็นมิติที่คะแนนต่ำสุด {score}/100 ควรพูดคุยความคาดหวัง จังหวะเวลา และขอบเขตอย่างชัดเจน",
"action_comm":"นัดคุยกันสัปดาห์ละ 20 นาที โดยแชร์ข้อสังเกต ความรู้สึก คำขอ และการกระทำที่ตกลงร่วมกัน",
"action_emotion":"ก่อนแก้ปัญหา ให้บอกชื่ออารมณ์และถามว่าอีกฝ่ายต้องการการสนับสนุนแบบใด",
"action_conflict":"เมื่อความตึงเครียดสูงขึ้น ให้หยุดอย่างน้อย 20 นาทีและกำหนดเวลาที่จะกลับมาคุย",
"action_stability":"ทำให้หน้าที่ เงิน เวลา ความผูกพัน และสิทธิในการตัดสินใจชัดเจน",
"limit":"คะแนนขึ้นอยู่กับข้อมูลเกิดที่ให้ไว้ หากไม่มีเวลาเกิดหรือพิกัด ความแม่นยำของลัคนา เรือน Zi Wei และแกน Composite จะลดลง",
"disclaimer":"รายงานนี้ช่วยการสะท้อนและการสื่อสาร ไม่ได้ตัดสินว่าความสัมพันธ์ต้องสำเร็จหรือล้มเหลว และไม่ใช่คำแนะนำทางการแพทย์ กฎหมาย หรือการเงิน",
},
"es": {
"summary":"Puntuación global {overall}/100. La dimensión más fuerte es {best} ({best_score}/100); el área principal de desarrollo es {weakest} ({weak_score}/100). La intensidad del conflicto es {conflict}/100.",
"strength":"{name} es una fortaleza relativa con {score}/100. Protégela mediante conductas consistentes, no mediante suposiciones.",
"challenge":"{name} es el área con menor puntuación ({score}/100). Hablen de expectativas, tiempos y límites concretos.",
"action_comm":"Realicen una revisión semanal de 20 minutos: una observación, una emoción, una petición y una acción acordada.",
"action_emotion":"Antes de resolver el problema, nombren la emoción y confirmen qué tipo de apoyo se necesita.",
"action_conflict":"Cuando aumente la tensión, hagan una pausa de al menos 20 minutos y acuerden una hora concreta para retomar.",
"action_stability":"Aclaren responsabilidades, dinero, tiempo, exclusividad y derechos de decisión.",
"limit":"Las puntuaciones dependen de los datos de nacimiento. Sin hora o coordenadas disminuye la precisión del Ascendente, casas, tiempos de Zi Wei y ángulos Composite.",
"disclaimer":"Este informe apoya la reflexión y la comunicación. No determina el éxito o fracaso de la relación y no constituye asesoramiento médico, legal ni financiero.",
},
"ar": {
"summary":"الدرجة العامة {overall}/100. أقوى بُعد هو {best} ({best_score}/100)، وأهم مجال للتطوير هو {weakest} ({weak_score}/100). شدة الصراع {conflict}/100.",
"strength":"يمثل {name} نقطة قوة نسبية بدرجة {score}/100. حافظا عليها بالسلوك المستمر لا بالافتراضات.",
"challenge":"يمثل {name} أقل الأبعاد بدرجة {score}/100. ناقشا التوقعات والتوقيت والحدود بصورة محددة.",
"action_comm":"خصصا 20 دقيقة أسبوعيًا لمشاركة ملاحظة وشعور وطلب وخطوة متفق عليها.",
"action_emotion":"قبل حل المشكلة، سمّيا المشاعر وحددا نوع الدعم المطلوب.",
"action_conflict":"عند ارتفاع التوتر، خذا استراحة لا تقل عن 20 دقيقة وحددا وقتًا للعودة إلى الحوار.",
"action_stability":"وضّحا المسؤوليات والمال والوقت والالتزام وحقوق اتخاذ القرار.",
"limit":"تعتمد الدرجات على بيانات الميلاد. غياب الوقت أو الإحداثيات يقلل دقة الطالع والبيوت وتوقيت Zi Wei ومحاور Composite.",
"disclaimer":"يدعم هذا التقرير التأمل والتواصل ولا يحدد نجاح العلاقة أو فشلها، وليس نصيحة طبية أو قانونية أو مالية.",
},
"zh-TW": {
"summary":"整體分數 {overall}/100。最高優勢為 {best}（{best_score}/100），最需要經營的是 {weakest}（{weak_score}/100），衝突張力為 {conflict}/100。",
"strength":"{name} 目前為相對優勢（{score}/100），需以穩定行動維持，而非依賴默契假設。",
"challenge":"{name} 為最低分維度（{score}/100），建議具體討論期待、時間與界線。",
"action_comm":"每週安排 20 分鐘關係檢視：一項觀察、一個感受、一個請求與一個共同行動。",
"action_emotion":"進入解決問題前，先說出情緒並確認對方需要哪種支持。",
"action_conflict":"張力升高時至少暫停 20 分鐘，並約定明確的恢復對話時間。",
"action_stability":"明確說清楚責任、金錢、時間、承諾與決策權。",
"limit":"分數依賴輸入的出生資料；缺少出生時間或座標會降低上升、宮位、紫微時辰與 Composite 軸線的精度。",
"disclaimer":"本報告用於關係理解與溝通參考，不決定關係必然成功或失敗，亦不構成醫療、法律或投資建議。",
},
}


def normalize_language(language: str | None) -> str:
    value = str(language or "zh-TW")
    aliases = {"zh":"zh-TW","zh_tw":"zh-TW","English":"en","日本語":"ja","ไทย":"th","Español":"es","العربية":"ar"}
    value = aliases.get(value, value)
    return value if value in _LANGS else "zh-TW"


def _sanitize_compat_name(name: str) -> str:
    cleaned = _re.sub(r'[\U0001F000-\U0001FFFF\U00002600-\U000027FF\U00002B00-\U00002BFF\U0000FE00-\U0000FEFF]', '', name)
    cleaned = _re.sub(r'[\\/:*?"<>|\r\n\t]', '', cleaned)
    cleaned = _re.sub(r'[\s\-]+', '_', cleaned).strip('_')
    return (cleaned or "person")[:40]


def make_compat_filename(name_a: str, name_b: str, ext: str, rel_type: str = "") -> str:
    return f"relationship_report_{_sanitize_compat_name(name_a)}_{_sanitize_compat_name(name_b)}_{datetime.now():%Y%m%d}.{ext}"


def _dimensions(report: "CompatibilityReport", language: str):
    tx = _TXT[language]; sc = report.score_breakdown
    return [
        ("emotional", tx["emotional"], int(sc.emotional_score)),
        ("communication", tx["communication"], int(sc.communication_score)),
        ("attraction", tx["attraction"], int(sc.attraction_score)),
        ("stability", tx["stability"], int(sc.stability_score)),
        ("growth", tx["growth"], int(sc.growth_score)),
        ("collaboration", tx["collaboration"], int(sc.collaboration_score)),
    ]


def _localized_report_data(report: "CompatibilityReport", language: str):
    language = normalize_language(language)
    tx = _TXT[language]; generic = _GENERIC[language]
    dims = _dimensions(report, language)
    best = max(dims, key=lambda item:item[2]); weakest = min(dims, key=lambda item:item[2])
    summary = generic["summary"].format(
        overall=report.score_breakdown.overall_score,
        best=best[1], best_score=best[2], weakest=weakest[1], weak_score=weakest[2],
        conflict=report.score_breakdown.conflict_score,
    )
    strengths = [generic["strength"].format(name=name, score=score) for _,name,score in sorted(dims,key=lambda x:x[2],reverse=True)[:2]]
    challenges = [generic["challenge"].format(name=weakest[1], score=weakest[2])]
    actions = []
    if report.score_breakdown.communication_score < 65: actions.append(generic["action_comm"])
    if report.score_breakdown.emotional_score < 65: actions.append(generic["action_emotion"])
    if report.score_breakdown.conflict_score >= 60: actions.append(generic["action_conflict"])
    if report.score_breakdown.stability_score < 65: actions.append(generic["action_stability"])
    if not actions: actions = [generic["action_comm"], generic["action_emotion"]]
    return {
        "tx":tx, "summary":summary, "strengths":strengths, "challenges":challenges,
        "actions":actions, "limit":generic["limit"], "disclaimer":generic["disclaimer"],
        "dims":dims,
    }


def export_compat_to_markdown(report: "CompatibilityReport", language: str = "zh-TW") -> str:
    language = normalize_language(language)
    if language == "zh-TW" and getattr(report, "markdown_body", ""):
        return report.markdown_body
    d = _localized_report_data(report, language); tx=d["tx"]
    pa, pb = report.person_a_profile, report.person_b_profile
    lines = [
        "---", f"language: {language}", f"direction: {_META[language][1]}", "---", "",
        f"# {tx['title']}", "", f"## {pa.name} × {pb.name}", "",
        f"- **{tx['relationship_type']}**: {getattr(report.relationship_type,'value',report.relationship_type)}",
        f"- **{tx['person_a']}**: {pa.name} — {pa.birth_date} — {pa.birth_city}",
        f"- **{tx['person_b']}**: {pb.name} — {pb.birth_date} — {pb.birth_city}",
        f"- **{tx['created']}**: {report.created_at}", "",
        f"## {tx['overall']}", "", f"**{report.score_breakdown.overall_score}/100**", "",
        f"## {tx['precision']}", "", d["summary"], "",
        f"## {tx['dimension']}", "", f"| {tx['dimension']} | Score |", "|---|---:|",
    ]
    for _,name,score in d["dims"]:
        lines.append(f"| {name} | {score}/100 |")
    lines += ["", f"## {tx['strengths']}", ""]
    lines += [f"- {item}" for item in d["strengths"]]
    lines += ["", f"## {tx['challenges']}", ""]
    lines += [f"- {item}" for item in d["challenges"]]
    lines += ["", f"## {tx['actions']}", ""]
    lines += [f"{i}. {item}" for i,item in enumerate(d["actions"],1)]
    lines += ["", f"## {tx['limits']}", "", d["limit"], "", f"## {tx['disclaimer']}", "", d["disclaimer"]]
    return "\n".join(lines)


def export_compat_to_html(report: "CompatibilityReport", language: str = "zh-TW") -> str:
    language = normalize_language(language)
    from reports.html_exporter import _CSS
    markdown_text = export_compat_to_markdown(report, language=language)
    try:
        import markdown as md
        body = md.markdown(markdown_text, extensions=["tables","fenced_code"])
    except Exception:
        body = f"<pre>{escape(markdown_text)}</pre>"
    lang, direction = _META[language]
    return (
        "<!DOCTYPE html>\n"
        f'<html lang="{lang}" dir="{direction}">\n<head>\n<meta charset="UTF-8">\n'
        '<meta name="viewport" content="width=device-width, initial-scale=1.0">\n'
        f"<title>{escape(_TXT[language]['title'])}</title>\n<style>{_CSS}\n"
        '[dir="rtl"] body{direction:rtl;text-align:right}[dir="rtl"] table{direction:rtl}'
        'code,pre,.url,.path{direction:ltr;unicode-bidi:embed}</style>\n</head>\n<body>\n'
        + body + "\n</body>\n</html>"
    )


def _docx_available() -> bool:
    try:
        import docx
        return True
    except ImportError:
        return False


def export_compat_to_docx(report: "CompatibilityReport", language: str = "zh-TW") -> bytes:
    if not _docx_available():
        raise RuntimeError("python-docx is required for Word export.")
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    language = normalize_language(language)
    text = export_compat_to_markdown(report, language=language)
    doc = Document()
    for line in text.splitlines():
        if line.startswith("# "): p=doc.add_heading(line[2:],0)
        elif line.startswith("## "): p=doc.add_heading(line[3:],1)
        elif line.startswith("### "): p=doc.add_heading(line[4:],2)
        elif line.startswith("- "): p=doc.add_paragraph(line[2:],style="List Bullet")
        elif _re.match(r"^\d+\. ",line): p=doc.add_paragraph(_re.sub(r"^\d+\. ","",line),style="List Number")
        elif line.startswith("|"): p=doc.add_paragraph(line)
        elif line.strip() in ("---",""): continue
        else: p=doc.add_paragraph(line)
        if language == "ar":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ppr=p._p.get_or_add_pPr()
            if ppr.find(qn("w:bidi")) is None: ppr.append(OxmlElement("w:bidi"))
    buf=BytesIO(); doc.save(buf); return buf.getvalue()


def _pdf_available() -> bool:
    try:
        import weasyprint
        return True
    except (ImportError,OSError):
        return False


def export_compat_to_pdf(report: "CompatibilityReport", language: str = "zh-TW") -> bytes:
    if _pdf_available():
        from weasyprint import HTML
        return HTML(string=export_compat_to_html(report,language=language)).write_pdf()
    # Always-available fallback using ReportLab (already a project dependency).
    try:
        from reports.pdf_fallback import markdown_to_pdf_bytes
        return markdown_to_pdf_bytes(export_compat_to_markdown(report,language=language),language=language)
    except Exception as exc:
        raise RuntimeError("PDF export requires WeasyPrint or ReportLab.") from exc
