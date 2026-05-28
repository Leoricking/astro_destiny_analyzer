"""
Astro Destiny Analyzer — Word (docx) Exporter
Uses python-docx to produce a structured Word document with cover,
disclaimer, basic-info table, calc-mode table, and chapter body.
Falls back gracefully if python-docx is not installed.
"""
from core.models import FullReport
from reports.templates import render_report
from reports.utils import build_report_meta
from config import APP_VERSION, APP_NAME, BRAND_NAME, REPORT_WATERMARK


def _docx_available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


def _set_cjk_font(paragraph, font_name: str = "Microsoft JhengHei") -> None:
    """Set CJK (eastAsia) font on all runs in a paragraph without crashing."""
    try:
        from docx.oxml.ns import qn
        for run in paragraph.runs:
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                from lxml import etree
                rFonts = etree.SubElement(rPr, qn("w:rFonts"))
            rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


class DocxExporter:
    def is_available(self) -> bool:
        return _docx_available()

    def export(self, report: FullReport) -> bytes:
        """Export report as .docx bytes."""
        if not self.is_available():
            raise ImportError(
                "python-docx is required. Run: pip install python-docx"
            )
        import docx
        from docx.shared import Pt
        from io import BytesIO

        meta = build_report_meta(report)
        markdown_text = render_report(report, version=APP_VERSION)
        document = docx.Document()

        # ── Default Normal style ──────────────────────────────────────────────
        try:
            style_normal = document.styles["Normal"]
            style_normal.font.name = "Arial"
            style_normal.font.size = Pt(11)
        except Exception:
            pass

        # ── Cover ─────────────────────────────────────────────────────────────
        document.add_heading(BRAND_NAME, level=0)
        document.add_heading("命盤整合分析報告", level=1)

        cover_table = document.add_table(rows=6, cols=2)
        try:
            cover_table.style = "Table Grid"
        except Exception:
            pass
        cover_data = [
            ("姓名", meta["name"]),
            ("出生日期", meta["birth_date"]),
            ("出生時間", meta["birth_time"]),
            ("出生地", meta["location"]),
            ("產生時間", meta["created_at"]),
            ("系統版本", f"v{APP_VERSION}"),
        ]
        for i, (k, v) in enumerate(cover_data):
            cover_table.cell(i, 0).text = k
            cover_table.cell(i, 1).text = v
        document.add_page_break()

        # ── Disclaimer ────────────────────────────────────────────────────────
        document.add_heading("免責聲明", level=2)
        disc_p = document.add_paragraph(meta["disclaimer"])
        try:
            if "Quote" in document.styles:
                disc_p.style = document.styles["Quote"]
        except Exception:
            pass
        _set_cjk_font(disc_p)
        document.add_paragraph()

        # ── Basic Info Table ──────────────────────────────────────────────────
        document.add_heading("基本資料", level=2)
        info_table = document.add_table(rows=8, cols=2)
        try:
            info_table.style = "Table Grid"
        except Exception:
            pass
        info_data = [
            ("姓名", meta["name"]),
            ("性別", meta["gender"]),
            ("出生日期", meta["birth_date"]),
            ("出生時間", meta["birth_time"]),
            ("出生地", meta["location"]),
            ("血型", meta["blood_type"]),
            ("分析主題", meta["themes"]),
            ("報告長度", meta["report_length"]),
        ]
        for i, (k, v) in enumerate(info_data):
            info_table.cell(i, 0).text = k
            info_table.cell(i, 1).text = v

        # ── Calc Mode Table ───────────────────────────────────────────────────
        document.add_heading("計算模式摘要", level=2)
        calc_table = document.add_table(rows=6, cols=3)
        try:
            calc_table.style = "Table Grid"
        except Exception:
            pass
        calc_table.cell(0, 0).text = "系統"
        calc_table.cell(0, 1).text = "計算模式"
        calc_table.cell(0, 2).text = "備注"
        calc_data = [
            ("西洋占星", meta["western_mode"], meta["western_note"]),
            ("八字", meta["bazi_mode"], meta["bazi_note"]),
            ("紫微", meta["ziwei_mode"], meta["ziwei_note"]),
            ("紫微輔星", "—", meta["ziwei_aux_note"]),
            ("紫微大限", meta["daxian_accuracy"], "—"),
        ]
        for i, (s, m, n) in enumerate(calc_data):
            calc_table.cell(i + 1, 0).text = s
            calc_table.cell(i + 1, 1).text = m
            calc_table.cell(i + 1, 2).text = n
        document.add_page_break()

        # ── Parse Markdown body into Word sections ────────────────────────────
        for line in markdown_text.splitlines():
            if line.startswith("# "):
                document.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                document.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                document.add_heading(line[4:].strip(), level=3)
            elif line.startswith("#### "):
                document.add_heading(line[5:].strip(), level=4)
            elif line.startswith("---"):
                document.add_paragraph("─" * 40)
            elif line.startswith("| "):
                # Table rows rendered as bullet list items (full table
                # parsing requires more complex logic)
                document.add_paragraph(line, style="List Bullet")
            elif line.startswith("- ") or line.startswith("* "):
                p = document.add_paragraph(line[2:].strip(), style="List Bullet")
                _set_cjk_font(p)
            elif line.startswith("> "):
                p = document.add_paragraph(line[2:].strip())
                try:
                    if "Quote" in document.styles:
                        p.style = document.styles["Quote"]
                except Exception:
                    pass
                _set_cjk_font(p)
            elif line.strip():
                p = document.add_paragraph(line.strip())
                _set_cjk_font(p)

        # ── Brand watermark footer ────────────────────────────────────────────
        document.add_paragraph()
        wm_p = document.add_paragraph(f"{REPORT_WATERMARK} · v{APP_VERSION}")
        try:
            wm_p.style = document.styles["Caption"]
        except Exception:
            pass

        # ── Save to bytes ─────────────────────────────────────────────────────
        buf = BytesIO()
        document.save(buf)
        return buf.getvalue()

    def save(self, report: FullReport, path: str) -> None:
        """Save .docx file to path."""
        data = self.export(report)
        with open(path, "wb") as f:
            f.write(data)
